import locale
import pandas as pd
from docx import Document
from address_parser import AddressParser

class DocxCreator:
    """Create a docx file from given addresses"""
    __template = 'addresses_template.docx'
    __style_normal = 'Normal'
    __style_committee = 'Heading 3'
    __style_function = 'Heading 3'

    def __init__(self, addresses, file_name):
        self.__file_name = file_name
        self.__addresses = addresses
        self.__document = Document(self.__template)

    def create(self):
        for committee, addresses in self.__addresses.groupby(AddressParser.COL_COMMITTEE, sort = False):
            self.__document.add_paragraph('', self.__style_normal)
            self.__document.add_paragraph(committee, self.__style_committee)
            for row_count, row in addresses.iterrows():
                self.__document.add_paragraph('', self.__style_normal)
                function = row[AddressParser.COL_FUNCTION]
                if not pd.isna(function):
                    self.__document.add_paragraph(row[AddressParser.COL_FUNCTION], self.__style_function)
                person_details = '{0} {1}\n{2}\n{3} {4}\n{5}'.format(
                        row[AddressParser.COL_FIRST_NAME],
                        row[AddressParser.COL_LAST_NAME],
                        row[AddressParser.COL_ADDRESS],
                        row[AddressParser.COL_ZIP],
                        row[AddressParser.COL_TOWN],
                        self.__get_phone_numbers(row)
                )
                self.__document.add_paragraph(person_details, self.__style_normal)
        self.__document.save(self.__file_name)

    def __get_phone_numbers(self, row):
        numbers = []
        number_private = row[AddressParser.COL_PHONE_PRIVATE]
        number_business = row[AddressParser.COL_PHONE_BUSINESS]
        number_mobile = row[AddressParser.COL_PHONE_MOBILE]
        if not pd.isna(number_private):
            numbers.append(number_private)
        if not pd.isna(number_business):
            numbers.append(number_business)
        if not pd.isna(number_mobile):
            numbers.append(number_mobile)
        return ', '.join(numbers)
