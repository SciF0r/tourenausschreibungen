import locale
import pandas as pd
import re
from docx import Document
from tour_parser import TourParser

class DocxCreatorRoteKarte:
    """Create a docx file from given tours"""
    __template = 'tourenausschreibungen_template.docx'
    __style = 'Standard_Ausschreibungen'
    __tours_sektion = []
    __tours_fabe = []
    __tours_kibe = []
    __tours_jo = []
    __tours_seniorenA = []
    __tours_seniorenB = []
    __tours_seniorenC = []
    __tours_alle = []
    __tours_versammlungen = []

    def __init__(self, tours, file_name):
        self.__file_name = file_name
        self.__tours = tours
        self.__document = Document(self.__template)
        locale.setlocale(locale.LC_ALL, 'de_CH')

    def create(self):
        for row_count, row in self.__tours.iterrows():
            groups = row[TourParser.COL_GROUP]
            if 'Sektion' in groups:
                self.__add_tour(row, self.__tours_sektion)
            if 'Familienbergsteigen' in groups:
                self.__add_tour(row, self.__tours_fabe)
            if 'Kinderbergsteigen' in groups:
                self.__add_tour(row, self.__tours_kibe)
            if 'Jugendorganisation' in groups:
                self.__add_tour(row, self.__tours_jo)
            if 'SeniorenA' in groups:
                self.__add_tour(row, self.__tours_seniorenA)
            if 'SeniorenB' in groups:
                self.__add_tour(row, self.__tours_seniorenB)
            if 'SeniorenC' in groups:
                self.__add_tour(row, self.__tours_seniorenC)
            if 'Alle' in groups:
                self.__add_tour(row, self.__tours_alle)
            if 'Versammlungen' in groups:
                self.__add_tour(row, self.__tours_versammlungen)
        self.__write_document()

    def __add_tour(self, row, tours_group):
        tour = []
        status = row[TourParser.COL_STATUS]
        if status == TourParser.STATUS_CANCELLED:
            self.__add_activity(row, tour)
            self.__add_guide(row, tour)
            self.__add_requirements(row, tour)
        elif status == TourParser.STATUS_FULL:
            self.__add_activity(row, tour)
            self.__add_guide(row, tour)
            self.__add_requirements(row, tour)
            self.__add_meeting_point(row, tour)
        else:
            self.__add_activity(row, tour)
            self.__add_guide(row, tour)
            self.__add_requirements(row, tour)
            self.__add_metrics(row, tour)
            self.__add_route(row, tour)
            self.__add_hospitality(row, tour)
            self.__add_costs(row, tour)
            self.__add_maps(row, tour)
            self.__add_execution(row, tour)
            self.__add_meeting_point(row, tour)
            self.__add_description(row, tour)
            self.__add_additional_information(row, tour)
            self.__add_gear(row, tour)
            self.__add_registration(row, tour)
        tours_group.append(tour)

    def __add_activity(self, row, tour):
        activity = row[TourParser.COL_ACTIVITY]
        status = row[TourParser.COL_STATUS]
        if status == TourParser.STATUS_CANCELLED:
            activity = '{0}, abgesagt'.format(activity)
        if status == TourParser.STATUS_FULL:
            activity = '{0}, ausgebucht'.format(activity)
        tour.append((self.__get_date(row), activity))

    def __add_guide(self, row, tour):
        guide = re.sub('(.*?), .*?, (Telefon|Mobile.*)', '\\1, \\2', row[TourParser.COL_GUIDE])
        tour.append((row[TourParser.COL_TOUR_TYPE_LONG], guide))

    def __add_requirements(self, row, tour):
        requirements = self.__get_requirements(row)
        if requirements:
            tour.append(('Anforderungen', requirements))

    def __add_metrics(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_METRICS, 'Auf-/Abstieg, MZ')

    def __add_route(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_ROUTE, 'Reiseroute')

    def __add_hospitality(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_HOSPITALITY, 'Unterk./Verpfl.')

    def __add_costs(self, row, tour):
        costs = self.__get_costs(row)
        if costs:
            tour.append(('Kosten', costs))

    def __add_maps(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_MAPS, 'Karten')

    def __add_execution(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_EXECUTION, 'Durchführung')

    def __add_meeting_point(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_MEETING_POINT, 'Treffpunkt')

    def __add_description(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_DESCRIPTION, 'Route / Details')

    def __add_additional_information(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_ADDITIONAL_INFORMATION, 'Zusatzinfo')

    def __add_gear(self, row, tour):
        self.__append_if_not_empty(row, tour, TourParser.COL_GEAR, 'Ausrüstung')

    def __add_registration(self, row, tour):
        registration = self.__get_registration(row)
        if registration:
            tour.append(('Anmeldung', registration))

    def __get_date(self, row):
        start_date = row[TourParser.COL_START_DATE]
        end_date = row[TourParser.COL_END_DATE]
        if pd.isna(start_date):
            return ''
        elif pd.isna(end_date):
            return start_date.strftime('%d.%m.%Y')
        else:
            return '{0}-{1}'.format(start_date.strftime('%d'), end_date.strftime('%d.%m.%y'))

    def __get_requirements(self, row):
        conditions = []
        cond_tech = row[TourParser.COL_TECH_REQ]
        if not pd.isna(cond_tech):
            conditions.append(cond_tech)
        cond_req = row[TourParser.COL_COND_REQ]
        if not pd.isna(cond_req):
            conditions.append(cond_req)
        return ', '.join(conditions)

    def __get_costs(self, row):
        costs = row[TourParser.COL_COSTS]
        costs_long = row[TourParser.COL_COSTS_LONG]
        costs_string = ''
        if costs == 0:
            return None
        elif pd.isna(costs_long):
            return 'Fr. {0}.--'.format(costs)
        else:
            return 'Fr. {0}.-- {1}'.format(costs, costs_long)

    def __get_registration(self, row):
        registration_means = row[TourParser.COL_REGISTRATION_MEANS]
        registration_start_date = row[TourParser.COL_REGISTRATION_START_DATE]
        registration_end_date = row[TourParser.COL_REGISTRATION_END_DATE]

        registration_fields = []
        if not pd.isna(registration_means):
            registration_fields.append(registration_means)
        if not pd.isna(registration_start_date):
            registration_fields.append('Anmeldestart {0}'.format(registration_start_date.strftime('%d.%m.%Y')))
        if not pd.isna(registration_end_date):
            registration_fields.append('Anmeldeschluss {0}'.format(registration_end_date.strftime('%d.%m.%Y')))
        return ', '.join(registration_fields)

    def __append_if_not_empty(self, row, tour, field, key):
        value = row[field]
        if not pd.isna(value):
            tour.append((key, value))

    def __write_document(self):
        self.__write_tours('Sektionstouren', self.__tours_sektion)
        self.__write_tours('Familienbergsteigen FaBe', self.__tours_fabe)
        self.__write_tours('Kinderbergsteigen KiBe', self.__tours_kibe)
        generic_registration_jo = 'Die Tourendetails erscheinen jeweils auf unserer Homepage (www.jo-aarau.ch) oder auf SAC Aarau-Tourenseite (www.sac-aarau.ch/touren.html), welche sich immer auf dem aktuellsten Stand befinden. Wir bitten euch, die Touren wenn möglich direkt über Drop-Tours anzumelden. Den Link dazu findet ihr auf unserer Homepage. Für Fragen oder Anmeldung zur JO-Aarau Mitgliedschaft bitten wir euch, euch direkt bei Danny Senn telefonisch (078 715 60 88) oder via E-Mail (info@jo-aarau.ch) zu melden.'
        self.__write_tours('Jugendorganisation JO', self.__tours_jo, generic_registration_jo)
        generic_registration_seniorenA = 'www.sac-aarau.ch oder Dienstag, 17-18 Uhr, telefonisch an die jeweilige Tourenleitung'
        self.__write_tours('Seniorengruppe A', self.__tours_seniorenA, generic_registration_seniorenA)
        generic_registration_seniorenB = 'Mail oder telefonisch am Dienstagabend 17-18 Uhr, beim jeweiligen Tourenleiter'
        self.__write_tours('Seniorengruppe B', self.__tours_seniorenB, generic_registration_seniorenB)
        generic_registration_seniorenC = 'Mail oder telefonisch am Mittwochabend 17-18 Uhr, beim jeweiligen Tourenleiter'
        self.__write_tours('Seniorengruppe C', self.__tours_seniorenC, generic_registration_seniorenC)
        self.__write_tours('Für alle Mitglieder SAC Aarau', self.__tours_alle)
        self.__write_tours('Versammlungen und Vorträge', self.__tours_versammlungen)
        print('Writing tours to {0}...'.format(self.__file_name))
        self.__document.save(self.__file_name)
        print('Done.')

    def __write_tours(self, group_title, tours, generic_registration = None):
        print('Adding {0} tours for {1}'.format(len(tours), group_title))
        if len(tours) == 0:
            return
        self.__document.add_heading(group_title, 2)
        if generic_registration:
            print('Adding generic registration text: {0}'.format(generic_registration))
            self.__document.add_heading('Anmeldung\t{0}'.format(generic_registration), 3)
            self.__document.add_paragraph('', self.__style)
        for tour in tours:
            left, right = tour.pop(0)
            self.__document.add_heading('{0}\t{1}'.format(left, right), 3)
            for left, right in tour:
                self.__document.add_paragraph('{0}\t{1}'.format(left, right), self.__style)
            self.__document.add_paragraph('', self.__style)
