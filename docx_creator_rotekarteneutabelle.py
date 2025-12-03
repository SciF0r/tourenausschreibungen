import locale
import pandas as pd
from docx import Document
from tour_parser import TourParser

class DocxCreatorRoteKarteNeuTabelle:
    """Create a docx file from given tours"""
    __template = 'jahresprogramm_template.docx'
    __style = 'Normal'

    def __init__(self, tours, file_name):
        self.__file_name = file_name
        self.__tours = tours
        self.__document_section = Document(self.__template)
        self.__table_section = self.__get_table_with_header(self.__document_section)
        self.__document_fabe = Document(self.__template)
        self.__table_fabe = self.__get_table_with_header(self.__document_fabe)
        self.__document_kibe = Document(self.__template)
        self.__table_kibe = self.__get_table_with_header(self.__document_kibe)
        self.__document_jo = Document(self.__template)
        self.__table_jo = self.__get_table_with_header(self.__document_jo)
        self.__document_all = Document(self.__template)
        self.__table_all = self.__get_table_with_header(self.__document_all)
        self.__document_events = Document(self.__template)
        self.__table_events = self.__get_table_with_header(self.__document_events)
        locale.setlocale(locale.LC_ALL, 'de_CH')

    def create(self):
        for row_count, row in self.__tours.iterrows():
            groups = row[TourParser.COL_GROUP]
            if 'Versammlungen' in groups:
                self.__add_line(row, self.__table_events)
            if 'Alle' in groups:
                self.__add_line(row, self.__table_all)
            if 'Sektion' in groups:
                self.__add_line(row, self.__table_section)
            if 'Familienbergsteigen' in groups:
                self.__add_line(row, self.__table_fabe)
            if 'Kinderbergsteigen' in groups:
                self.__add_line(row, self.__table_kibe)
            if 'Jugendorganisation' in groups:
                self.__add_line(row, self.__table_jo)
        self.__document_section.save(self.__file_name.format('sektion'))
        self.__document_fabe.save(self.__file_name.format('fabe'))
        self.__document_kibe.save(self.__file_name.format('kibe'))
        self.__document_jo.save(self.__file_name.format('jo'))
        self.__document_all.save(self.__file_name.format('alle'))
        self.__document_events.save(self.__file_name.format('anlässe'))

    def __get_table_with_header(self, document):
        table = document.add_table(1, 5)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Beginn'
        heading_cells[1].text = 'Dauer'
        heading_cells[2].text = 'Tour / Anlass'
        heading_cells[3].text = 'Typ / Kondition, Technik'
        heading_cells[4].text = 'Tourenleitung'
        return table

    def __add_line(self, row, table):
        cells = table.add_row().cells
        cells[0].text = self.__get_date(row)
        cells[1].text = self.__get_duration(row)
        cells[2].text = self.__get_activity(row)
        cells[3].text = '{0} / {1}'.format(self.__get_type(row), self.__get_requirements(row))
        cells[4].text = self.__get_guide(row)

    def __get_date(self, row):
        return row[TourParser.COL_START_DATE].strftime('%a, %d. %b')

    def __get_duration(self, row):
        duration = row[TourParser.COL_DURATION]
        duration = duration.replace(' Tage', '')
        duration = duration.replace(' Tag', '')
        return duration

    def __get_type(self, row):
        return row[TourParser.COL_TOUR_TYPE]

    def __get_requirements(self, row):
        conditions = []
        cond_req = row[TourParser.COL_COND_REQ]
        if not pd.isna(cond_req):
            conditions.append(cond_req)
        cond_tech = row[TourParser.COL_TECH_REQ]
        if not pd.isna(cond_tech):
            conditions.append(cond_tech)
        return ', '.join(conditions)

    def __get_activity(self, row):
        return row[TourParser.COL_ACTIVITY]
    
    def __get_guide(self, row):
        if pd.isna(row[TourParser.COL_FIRST_NAME]):
            return ''
        return '{0} {1}'.format(row[TourParser.COL_FIRST_NAME], row[TourParser.COL_LAST_NAME])

    def __is_whole_day(self, row):
        return 'Tag' in row[TourParser.COL_DURATION]
