from docx import Document

class DocxCreator:
    """Create a docx file from given tours"""
    __document = 'tourenausschreibungen_template.docx'
    __tour_line_style = 'Standard_Ausschreibungen'

    def __init__(self, tours, file_name):
        self.__file_name = file_name
        self.__tours = tours
        self.__document = Document(self.__document)

    def create(self):
        for tour_type, tours_of_type in self.__tours.items():
            self.__document.add_heading(tour_type, 2)
            for tour in tours_of_type:
                self.__add_tour(tour)
        self.__document.save(self.__file_name)

    def __add_tour(self, tour):
        left, right = tour.pop(0)
        self.__document.add_heading('{0}\t{1}'.format(left, right), 3)
        for left, right in tour:
            self.__add_tour_line('{0}\t{1}'.format(left, right))
        self.__add_tour_line('')

    def __add_tour_line(self, line):
        self.__document.add_paragraph(line, self.__tour_line_style)
