import locale
import pandas as pd
from docx import Document
from tour_parser import TourParser

class DocxCreatorJahresprogramm:
    """Create a docx file from given tours"""
    __template = 'jahresprogramm_template.docx'
    __style = 'Normal'

    def __init__(self, tours, file_name):
        self.__file_name = file_name
        self.__tours = tours
        self.__document_section = Document(self.__template)
        self.__document_fabe = Document(self.__template)
        self.__document_kibe = Document(self.__template)
        self.__document_jo = Document(self.__template)
        self.__document_all = Document(self.__template)
        self.__document_events = Document(self.__template)
        locale.setlocale(locale.LC_ALL, 'de_CH')

    def create(self):
        for row_count, row in self.__tours.iterrows():
            groups = row[TourParser.COL_GROUP]
            if 'Versammlungen' in groups:
                self.__add_line(row, self.__document_events)
            if 'Alle' in groups:
                self.__add_line(row, self.__document_all)
            if 'Sektion' in groups:
                self.__add_line(row, self.__document_section)
            if 'Familienbergsteigen' in groups:
                self.__add_line(row, self.__document_fabe)
            if 'Kinderbergsteigen' in groups:
                self.__add_line(row, self.__document_kibe)
            if 'Jugendorganisation' in groups:
                self.__add_line(row, self.__document_jo)
        self.__document_section.save(self.__file_name.format('sektion'))
        self.__document_fabe.save(self.__file_name.format('fabe'))
        self.__document_kibe.save(self.__file_name.format('kibe'))
        self.__document_jo.save(self.__file_name.format('jo'))
        self.__document_all.save(self.__file_name.format('alle'))
        self.__document_events.save(self.__file_name.format('anlässe'))

    def __add_line(self, row, document):
        if (self.__is_whole_day(row)):
            line = '{0}\t{1}\t{2}\t{3}\t{4}'.format(self.__get_date(row), self.__get_duration(row), self.__get_type(row), self.__get_requirements(row), self.__get_info(row))
        else:
            line = '{0}\t{1}, {2}\t{3}\t{4}'.format(self.__get_date(row), self.__get_duration(row), self.__get_type(row), self.__get_requirements(row), self.__get_info(row))
        document.add_paragraph(line, self.__style)

    def __get_date(self, row):
        return row[TourParser.COL_START_DATE].strftime('%a, %d. %b')

    def __get_duration(self, row):
        duration = row[TourParser.COL_DURATION]
        duration = duration.replace(' Tage', '')
        duration = duration.replace(' Tag', '')
        return duration

    def __get_type(self, row):
        return row[TourParser.COL_TOUR_TYPE]

    def __get_requirements(self, row):
        conditions = []
        cond_req = row[TourParser.COL_COND_REQ]
        if not pd.isna(cond_req):
            conditions.append(cond_req)
        cond_tech = row[TourParser.COL_TECH_REQ]
        if not pd.isna(cond_tech):
            conditions.append(cond_tech)
        return ', '.join(conditions)

    def __get_info(self, row):
        return '{0}\n{1} {2}'.format(row[TourParser.COL_ACTIVITY], row[TourParser.COL_FIRST_NAME], row[TourParser.COL_LAST_NAME])

    def __is_whole_day(self, row):
        return 'Tag' in row[TourParser.COL_DURATION]
